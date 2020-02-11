"""
Wrapper to the Bootstrap Layout library
"""


from epyk.core.html import Html
from epyk.core.html import HtmlSelect
from epyk.core.html.options import OptPanel

#
from epyk.core.js import JsUtils
from epyk.core.js.html import JsHtmlPanels

from epyk.core.css import Defaults

# The list of CSS classes
from epyk.core.css.styles.classes import CssStyle
from epyk.core.css.styles import GrpClsContainer


class Panel(Html.Html):
  name, category, callFnc = 'Panel', 'Layouts', 'panel'

  def __init__(self, report, htmlObj, title, color, width, height, htmlCode, helper, profile):
    if isinstance(htmlObj, list) and htmlObj:
      for obj in htmlObj:
        if hasattr(obj, 'inReport'):
          obj.inReport = False
    elif htmlObj is not None and hasattr(htmlObj, 'inReport'):
      htmlObj.inReport = False # Has to be defined here otherwise it is set to late
    component = []
    if title is not None:
      self.title = report.ui.title(title)
      self.title.inReport = False
      component.append(self.title)
    container = report.ui.div(htmlObj)
    container.inReport = False
    component.append(container)
    super(Panel, self).__init__(report, component, code=htmlCode, profile=profile,
                                css_attrs={"color": color, "width": width, "height": height})
    container.set_attrs(name="name", value="panel_%s" % self.htmlId)

  @property
  def style(self):
    if self._styleObj is None:
      self._styleObj = GrpClsContainer.ClassDiv(self)
    return self._styleObj

  def __add__(self, htmlObj):
    """ Add items to a container """
    htmlObj.inReport = False # Has to be defined here otherwise it is set to late
    self.val.append(htmlObj)
    return self

  @property
  def dom(self):
    """
    Javascript Functions

    Return all the Javascript functions defined for an HTML Component.
    Those functions will use plain javascript by default.

    :return: A Javascript Dom object
    :rtype: JsHtmlPanels.JsHtmlPanel
    """
    if self._dom is None:
      self._dom = JsHtmlPanels.JsHtmlPanel(self, report=self._report)
    return self._dom

  def __str__(self):
    str_div = "".join([v.html() if hasattr(v, 'html') else str(v) for v in self.val])
    return "<div %s>%s</div>%s" % (self.get_attrs(pyClassNames=self.style.get_classes()), str_div, self.helper)


class PanelSplit(Html.Html):
  __reqJs, __reqCss = ['jqueryui'], ['jqueryui']
  name, category, callFnc = 'Panel Split', 'Layouts', 'panelsplit'

  def __init__(self, report, width, height, left_width, left_obj, right_obj, resizable, helper, profile):
    super(PanelSplit, self).__init__(report, None, width=width[0], widthUnit=width[1], height=height[0],
                                     heightUnit=height[1], profile=profile)
    self.left_width, self.htmlMaps, self.resizable = left_width, {}, resizable
    if left_obj is not None:
      self.left(left_obj)
    if right_obj is not None:
      self.right(right_obj)
    self.css_left = {'flex': '0 0 auto', 'padding': '5px', 'min-width': '100px', 'width': self.left_width,
                     'white-space': 'nowrap'}
    self.css_right = {'flex': '0 1 auto', 'padding': '5px', 'width': '100%', 'background': self._report.theme.greys[0],
                     'border-left': '3px solid %s' % self._report.theme.success[1]}
    self.css({'display': 'flex', 'flex-direction': 'row', 'overflow': 'hidden', 'xtouch-action': 'none'})

  def left(self, html_obj):
    """
    Add the left component to the panel

    :param html_obj:
    """
    html_obj.inReport = False
    self.html_left = html_obj
    return self

  def right(self, html_obj):
    """
    Add the right component to the panel

    :param html_obj:
    """
    html_obj.inReport = False
    self.html_right = html_obj
    return self

  def __str__(self):
    self._report._props.setdefault('js', {}).setdefault("builders", []).extend([
      '$("#%(htmlId)s_left").resizable({handleSelector: ".splitter", resizeHeight: false});' % {'htmlId': self.htmlId},
      '$("#%(htmlId)s_right").resizable({handleSelector: ".splitter-horizontal", resizeWidth: true})' % {
        'htmlId': self.htmlId}])
    return '''
      <div %(attrs)s>
        <div style="%(css_left)s" id="%(htmlId)s_left" class="panel-left">%(left)s</div>
        <div style="%(css_right)s" id="%(htmlId)s_right" class="panel-right">%(right)s</div>
      </div>
      ''' % {"attrs": self.get_attrs(pyClassNames=self.style.get_classes()), "htmlId": self.htmlId, 'left': self.html_left.html(),
             'right': self.html_right.html(), 'css_left': CssStyle.get_style(self.css_left), 'css_right': CssStyle.get_style(self.css_right)}


class PanelSlide(Panel):
  __reqCss, __reqJs = ['font-awesome'], ['font-awesome']
  name, category, callFnc = 'Slide Panel', 'Panels', 'slide'

  def __init__(self, report, htmlObj, title, color, width, height, htmlCode, helper, options, profile):
    super(PanelSlide, self).__init__(report, htmlObj, title, color, width, height, htmlCode, helper, profile)
    self.title._vals = "<i style='float:left;margin:4px 5px 0 0' name='icon_%s' class='fas fa-caret-down'></i>%s" % (self.htmlId, self.title._vals)
    self.title.click([
      report.js.getElementsByName("panel_%s" % self.htmlId).first.toggle(),
      report.js.getElementsByName("icon_%s" % self.htmlId).first.toggleClass("fa-caret-up")])
    self.title.css({"cursor": 'pointer', "padding": "0 2px 0 0"})


class Div(Html.Html):
  name, category, callFnc = 'Simple Container', 'Layouts', 'div'

  def __init__(self, report, htmlObj, label, color, width, icon, height, editable, align, padding, htmlCode, tag,
               helper, profile):
    if isinstance(htmlObj, list) and htmlObj:
      newHtmlObj = []
      for obj in htmlObj:
        if isinstance(obj, list) and obj:
          newHtmlObj.append(report.ui.div(obj, label, color, width, icon, height, editable, align, padding,
                                          htmlCode, tag, helper, profile))
        else:
          newHtmlObj.append(obj)
        if hasattr(newHtmlObj[-1], 'inReport'):
          newHtmlObj[-1].inReport = False
          if newHtmlObj[-1].css("display") != 'None':
            newHtmlObj[-1].css({"display": 'inline-block'})
      htmlObj = newHtmlObj
    elif htmlObj is not None and hasattr(htmlObj, 'inReport'):
      htmlObj.inReport = False # Has to be defined here otherwise it is set to late
    super(Div, self).__init__(report, htmlObj, code=htmlCode, css_attrs={"color": color, "width": width, "height": height},
                              profile=profile)
    self.htmlMaps, self.tag = {}, tag
    # Add the component predefined elements
    self.add_icon(icon)
    self.add_label(label)
    self.add_helper(helper)

    self.css({'text-align': align})
    if padding is not None:
      self.css('padding', '%s' % padding)
    if editable:
      self.set_attrs(name='contenteditable', value="true")
      self.css('overflow', 'auto')

  def __add__(self, htmlObj):
    """ Add items to a container """
    htmlObj.inReport = False # Has to be defined here otherwise it is set to late
    if not isinstance(self.val, list):
      self.val = [self.val]
    self.val.append(htmlObj)
    return self

  def __getitem__(self, i):
    return self.val[i]

  @property
  def style(self):
    if self._styleObj is None:
      self._styleObj = GrpClsContainer.ClassDiv(self)
    return self._styleObj

  def build(self, data=None, options=None, profile=False):
    if isinstance(data, dict):
      # check if there is no nested HTML components in the data
      js_data = "{%s}" % ",".join(["%s: %s" % (k, JsUtils.jsConvertData(v, None)) for k, v in data.items()])
    else:
      js_data = JsUtils.jsConvertData(data, None)
    options, js_options = options or {}, []
    for k, v in options.items():
      if isinstance(v, dict):
        row = ["%s: %s" % (s_k, JsUtils.jsConvertData(s_v, None)) for s_k, s_v in v.items()]
        js_options.append("%s: {%s}" % (k, ", ".join(row)))
      else:
        js_options.append("%s: %s" % (k, JsUtils.jsConvertData(v, None)))
    return "%s.innerHTML = %s" % (self.dom.varId, js_data) #, "{%s}" % ",".join(js_options))

  def __str__(self):
    str_div = "".join([v.html() if hasattr(v, 'html') else str(v) for v in self.val])
    return "<div %s>%s</div>%s" % (self.get_attrs(pyClassNames=self.style.get_classes()), str_div, self.helper)

  # -----------------------------------------------------------------------------------------
  #                                    EXPORT OPTIONS
  # -----------------------------------------------------------------------------------------
  def to_word(self, document):
    if isinstance(self.vals, list):
      for val in self.vals:
        if hasattr(val, 'inReport'):
          val.to_word(document)
    else:
      if hasattr(self.vals, 'inReport'):
        self.vals.to_word(document)


class Table(Html.Html):
  name, category, callFnc = 'Row', 'Layouts', 'row'

  def __init__(self, report, htmlObjs, width, height, data, align, valign, colsWith, closable, resizable, titles,
               helper, profile):
    if data is not None:
      # Load the different HTML components from a static list
      # This mode will automatically add the inReport to the new components
      htmlObjs = []
      for component in data:
        fnc = getattr(report, component['htmlComponent'])
        parameters = dict(component)
        del parameters['htmlComponent']

        htmlObjs.append(fnc(**parameters))
    self.colsWith, self.htmlMaps = [] if colsWith is None else colsWith, {}
    super(Table, self).__init__(report, [], width=width[0], widthUnit=width[1], height=height[0], heightUnit=height[1], profile=profile)
    for htmlObj in htmlObjs:
      self.__add__(htmlObj)
    self.align, self.valign, self.closable, self.resizable, self.titles = align, valign, closable, resizable, titles
    self.css({"margin-top": '5px', "padding": "5px 0 5px 0", "border-collapse": "collapse", 'table-layout': 'fixed'})

  def __add__(self, htmlObj):
    """ Add items to a container """
    htmlObj.inReport = False # Has to be defined here otherwise it is set to late
    self.val.append(htmlObj)
    self.htmlMaps[htmlObj.htmlId] = htmlObj
    return self

  def get(self, htmlCode):
    """
    Return the Html component in the parameter bar

    :param htmlCode: The htmlCode for the component as a String
    """
    return self.htmlMaps[htmlCode]

  def __str__(self):
    """ Return the HTML display of a split container"""
    items = ['<div style="width:100%%;display:block;"><table %s><tr>' % self.get_attrs(pyClassNames=self.pyStyle)]
    widths = {}
    if self.colsWith:
      for i, _ in enumerate(self.vals):
        widths[i] = 'width:%s' % self.colsWith[i]
    if self.closable:
      if self.titles:
        for i, htmlObj in enumerate(self.vals):
          onclickEvent = "ResizableRow(this, 'col_%s_%s')" % (self.htmlId, i)  # if self.colsWith else "$(\'#col_%s_%s\').children().toggle()" % ( self.htmlId, i)
          items.append('<th style="text-align:left;padding:5px 0 5px 0;%s"><i onclick="%s" style="cursor:pointer;" class="far fa-minus-square"></i>&nbsp;<p style="color:%s;display:inline">%s</p></th>' % (widths.get(i, ''), onclickEvent, self._report.theme.danger[1], self.titles[i].upper()))
      else:
        for i, htmlObj in enumerate(self.vals):
          onclickEvent = "ResizableRow(this, 'col_%s_%s'))" % (self.htmlId, i)  # if self.colsWith else "$(\'#col_%s_%s\').children().toggle()" % ( self.htmlId, i)
          items.append( '<th style="text-align:left;%s"><i onclick="%s" style="cursor:pointer;" class="far fa-minus-square"></i></th>' % (
            widths.get(i, ''), onclickEvent))
      items.append('</tr><tr>')  # $(htmlId).parent().width()
      self.addGlobalFnc("ResizableRow(htmlId, targetId)", '''
         if ( $(htmlId).parent().data('size') == undefined) { $(htmlId).parent().data('size', $(this).parent().width() ) } ; 
         $('#' + targetId).children().toggle(); var styleDisplay = $('#' + targetId).children().css('display') ;
         if ( styleDisplay == 'block') { $(htmlId).parent().width($(htmlId).parent().data('size')) ; }
         else { $(htmlId).parent().width(10) ; } ''')

    for i, htmlObj in enumerate(self.val):
      extraStyle = 'padding:0 0 0 10px' if i != 0 else 'padding:0'
      items.append( '<td id="col_%s_%s" style="font-size:inherit;line-height:inherit;vertical-align:%s;text-align:%s;%s;%s">%s</td>' % (self.htmlId, i, self.valign, self.align, widths.get(i, ''), extraStyle, htmlObj.html()))
    items.append('</tr></table></div>')
    if self.resizable:
      self._report.jsImports.add('datatables-col-resizable')
      self._report.jsOnLoadFnc.add("%s.colResizable({ liveDrag:true });" % self.jqId)
    return "".join(items)


class Col(Html.Html):
  name, category, callFnc = 'Column', 'Layouts', 'col'

  def __init__(self, report, htmlObjs, position, width, height, align, helper, profile):
    self.position, self.htmlMaps, self.rows_css, self.row_css_dflt = position, {}, {}, {}
    super(Col, self).__init__(report, [], width=width[0], widthUnit=width[1], height=height[0], heightUnit=height[1], profile=profile)
    if htmlObjs is not None:
      for htmlObj in htmlObjs:
        self.__add__(htmlObj)
    if align == "center":
      self.css({'margin': 'auto', 'display': 'inline-block', 'text-align': 'center'})
    else:
      self.css({'display': 'inline-block'})

  def __add__(self, htmlObj):
    """ Add items to a container """
    if not hasattr(htmlObj, 'inReport'):
      # Add a text HTML internal object by default
      # todo: add options to this component to remove hard coded Css
      htmlObj = self._report.ui.div(htmlObj)
      htmlObj.style.addCls("CssDivOnHover")
    htmlObj.inReport = False # Has to be defined here otherwise it is set to late
    self.val.append(htmlObj)
    self.htmlMaps[htmlObj.htmlId] = htmlObj
    return self

  def get(self, htmlCode):
    """
    Return the Html component in the parameter bar

    :param htmlCode: The htmlCode for the component as a String
    """
    return self.htmlMaps[htmlCode]

  def set_css_row(self, css_attrs, row_id=None):
    """
    Set the CSS attributes for the row container

    :param css_attrs: The CSS attributes
    :param row_id: The row id for the special classes. None if it should be applied to all the rows

    :return: self to allow the chains
    """
    if row_id is None:
      self.row_css_dflt = css_attrs
    self.rows_css[row_id] = dict(self.row_css_dflt)
    self.rows_css[row_id].update(css_attrs)
    return self

  def __str__(self):
    self.css({"justify-content": self.position})
    rows = []
    for i, htmlObj in enumerate(self.val):
      css_style = "style='%s'" % ";".join(["%s:%s" % (k, v) for k, v in self.rows_css.get(i, self.row_css_dflt).items()])
      rows.append("<div %s>%s</div>" % (css_style, htmlObj.html()))
    return '<div %s>%s</div>' % (self.get_attrs(), "".join(rows))

  # -----------------------------------------------------------------------------------------
  #                                    EXPORT OPTIONS
  # -----------------------------------------------------------------------------------------
  def to_word(self, document):
    for i, htmlObj in enumerate(self.vals):
      htmlObj.to_word(document)

  def to_xls(self, workbook, worksheet, cursor):
    for i, htmlObj in enumerate(self.vals):
      try:
        htmlObj.to_xls(workbook, worksheet, cursor)
      except Exception as err:
        cell_format = workbook.add_format({'bold': True, 'font_color': 'red'})
        worksheet.write(cursor['row'], 0, str(err), cell_format)
        cursor['row'] += 2


class Row(Html.Html):
  name, category, callFnc = 'Column', 'Layouts', 'col'
  __reqCss, __reqJs = ['bootstrap'], ['bootstrap']

  def __init__(self, report, htmlObjs, position, width, height, align, helper, profile):
    self.position, self.htmlMaps = position, {}
    super(Row, self).__init__(report, [], css_attrs={"width": width, "height": height}, profile=profile)
    if htmlObjs is not None:
      for htmlObj in htmlObjs:
        self.__add__(htmlObj)
    self.attr["class"].add('row')

  def __add__(self, htmlObj):
    """ Add items to a container """
    if not hasattr(htmlObj, 'inReport'):
      # Add a text HTML internal object by default
      # todo: add options to this component to remove hard coded Css
      htmlObj = self._report.ui.div(htmlObj)
      htmlObj.attr["class"].add("CssDivOnHover")
    htmlObj.inReport = False # Has to be defined here otherwise it is set to late
    self.val.append(htmlObj)
    self.htmlMaps[htmlObj.htmlId] = htmlObj
    return self

  def get(self, htmlCode):
    """
    Return the Html component in the parameter bar

    :param htmlCode: The htmlCode for the component as a String
    """
    return self.htmlMaps[htmlCode]

  def __getitem__(self, i):
    return self.val[i]

  def __str__(self):
    self.css({"justify-content": self.position})
    rows = []
    for i, htmlObj in enumerate(self.val):
      htmlObj.attr["class"].add("col col-md-2")
      rows.append(htmlObj.html())
    return '<div %s>%s</div>' % (self.get_attrs(pyClassNames=self.style.get_classes()), "".join(rows))


class Grid(Html.Html):
  name, category, callFnc = 'Grid', 'Layouts', 'grid'
  __reqCss, __reqJs = ['bootstrap'], ['bootstrap']

  def __init__(self, report, htmlObjs, width, height, colsDim, colsAlign, noGlutters, align, helper, profile):
    super(Grid, self).__init__(report, [], width=width[0], widthUnit=width[1], height=height[0], heightUnit=height[1], profile=profile)
    self.css({'overflow-x': 'hidden', 'padding': 0})
    #self.attr["class"].add("container-fluid")
    self.rowsStyle, self.colsStyle, self.noGlutters = {}, {}, noGlutters
    if align == 'center':
      self.css({'margin': 'auto'})
    self.colsDim, self.htmlMaps, self.colsAlign = [], {}, []
    if colsDim is None:
      colsDim, currDim = [], 0
      for h in range(0, len(htmlObjs)-1):
        currDim += int(12 / len(htmlObjs))
        colsDim.append(int(12 / len(htmlObjs)))
      colsDim.append(12 - currDim)
    for i, htmlObj in enumerate(htmlObjs):
      self.__add__((htmlObj, colsDim[i]))
    if colsAlign is not None:
      self.colsAlign = colsAlign

  def __add__(self, htmlObjWithDim):
    """ Add items to a container """
    if isinstance(htmlObjWithDim, tuple):
      htmlObj, dim = htmlObjWithDim
    else:
      htmlObj, dim = htmlObjWithDim, 1
    self.htmlMaps[htmlObj.htmlId] = htmlObj
    htmlObj.inReport = False # Has to be defined here otherwise it is set to late
    self.val.append(htmlObj)
    self.colsDim.append(dim)
    self.colsAlign.append("left")
    return self

  def __getitem__(self, i):
    return self.val[i]

  @property
  def dom(self):
    """
    Javascript Functions

    Return all the Javascript functions defined for an HTML Component.
    Those functions will use plain javascript by default.

    :return: A Javascript Dom object
    :rtype: JsHtmlPanels.JsHtmlGrid
    """
    if self._dom is None:
      self._dom = JsHtmlPanels.JsHtmlGrid(self, report=self._report)
    return self._dom

  def get(self, htmlCode):
    """
    Return the Html component in the parameter bar

    :param htmlCode: The htmlCode for the component as a String
    """
    return self.htmlMaps[htmlCode]

  def resize(self):
    """
    For the resizing of the space for the containers.

    This will rescale based on the number of items and the fact that the max per row is 12
    It will update the colsDim list
    """
    max_size = int(12 / len(self.colsDim))
    self.colsDim = [max_size] * len(self.colsDim)
    return self

  def __str__(self):
    items = ['<div %s>' % self.get_attrs(pyClassNames=self.style.get_classes())]
    items.append('<div class="row%s">' % (' no-gutters' if self.noGlutters else ''))
    dim_row, row_index, col_per_obj = 0, 1, {}
    for i, htmlObj in enumerate(self.val):
      if dim_row == 12:
        items.append('</div><div class="row%s">' % (' no-gutters' if self.noGlutters else ''))
        dim_row = 0

      if isinstance(htmlObj, HtmlSelect.Select):
        htmlObj.container = "#%s" % self.htmlId # The container should be defined in this case to be visible
      htmlContent = htmlObj.html() # htmlObj.content() if isinstance(htmlObj, Col) else htmlObj.html()
      items.append('<div class="col col-md-%s text-%s">%s</div>' % (self.colsDim[i], self.colsAlign[i], htmlContent))
      dim_row += 1 if self.colsDim[i] == 'auto' else self.colsDim[i]
      col_per_obj[i] = self.colsDim[i]
      row_index += 1
      if dim_row > 12:
        raise Exception("BootStrap allow a max of 12 columns per Row")
    self._report.js.getVar("panel_dims_%s" % self.htmlId, col_per_obj)
    items.append('</div></div>')
    return "".join(items)

  # -----------------------------------------------------------------------------------------
  #                                    EXPORT OPTIONS
  # -----------------------------------------------------------------------------------------
  def to_word(self, document):
    for i, htmlObj in enumerate(self.vals):
      try:
        htmlObj.to_word(document)
      except Exception as err:
        from docx.shared import RGBColor

        errotTitle = document.add_heading().add_run("Error")
        errotTitle.font.color.rgb = RGBColor(255, 0, 0)
        errotTitle.font.italic = True
        errorParagraph = document.add_paragraph().add_run((str(err)))
        errorParagraph.font.color.rgb = RGBColor(255, 0, 0)
        errorParagraph.font.italic = True

  def to_xls(self, workbook, worksheet, cursor):
    for i, htmlObj in enumerate(self.vals):
      try:
        htmlObj.to_xls(workbook, worksheet, cursor)
      except Exception as err:
        cell_format = workbook.add_format({'bold': True, 'font_color': 'red'})
        worksheet.write(cursor['row'], 0, str(err), cell_format)
        cursor['row'] += 2

  # -----------------------------------------------------------------------------------------
  #                                    MARKDOWN SECTION
  # -----------------------------------------------------------------------------------------
  @staticmethod
  def matchMarkDown(val): return True if val == "[" else None

  @classmethod
  def convertMarkDown(cls, val, regExpResult, report=None):
    pass

  @classmethod
  def jsMarkDown(self, vals):
    return '''
      '''


class Tabs(Html.Html):
  name, category, callFnc = 'Tabs', 'Layouts', 'tabs'

  def __init__(self, report, color, size, width, height, htmlCode, helper, css_tab, options, profile):
    super(Tabs, self).__init__(report, "", code=htmlCode, width=width[0], widthUnit=width[1], height=height[0],
                               heightUnit=height[1], profile=profile)
    self.__panels, self.__panel_objs = [], {}
    self.tabs_name, self.panels_name = "button_%s" % self.htmlId, "panel_%s" % self.htmlId
    self.css_tab = css_tab
    self.options = options
    self.css_tab_clicked_dflt = {"border-bottom": "1px solid %s" % self._report.theme.success[1]}
    self.tabs_container = self._report.ui.div([])
    self.tabs_container.inReport = False

  def __getitem__(self, name):
    return self.__panel_objs[name]

  def panel(self, name):
    """

    :param name:

    :rtype: Div
    :return:
    """
    return self[name]["content"]

  def tab(self, name):
    """

    :param name:

    :rtype: Div
    :return:
    """
    return self[name]["tab"]

  def add_panel(self, name, div, icon=None, selected=False, css_tab=None, css_tab_clicked=None):
    """

    :param name:
    :param div:
    :param selected:
    :param css_tab:

    :return:
    """
    div.css({"display": 'none'})
    div.inReport = False
    div.set_attrs(name="name", value=self.panels_name)
    self.__panels.append(name)
    if icon is not None:
      tab = self._report.ui.div([
        self._report.ui.icon(icon).css({"display": 'block', "width": '100%',
                                        "font-size": '%spx' % (Defaults.Font.header_size + 4)}),
        name], width=("100", "px"))
    else:
      tab = self._report.ui.div(name, width=("100", "px"))
    css_tab = dict(self.css_tab)
    dflt_css_clicked, css_not_clicked = dict(self.css_tab_clicked_dflt), {}
    if css_tab is not None:
      css_tab.update(css_tab)
    if css_tab_clicked is not None:
      dflt_css_clicked = css_tab_clicked
    for key, val in dflt_css_clicked.items():
      if key in css_tab:
        css_not_clicked[key] = css_tab[key]
      else:
        css_not_clicked[key] = 'none'
    tab.css(css_tab).css({"padding": '5px 0'})
    tab.set_attrs(name="name", value=self.tabs_name)
    tab_container = self._report.ui.div(tab, width=("100", "px"))
    tab_container.inReport = False
    tab_container.css({'display': 'inline-block'})
    css_cls_name = None
    if self.options.get("tab_class") is not None:
      tab_container.defined.add(self.options.get("tab_class"), toMain=False)
      css_cls_name = CssStyle.cssName(self.options.get("tab_class"))
    tab.click([
      self._report.js.getElementsByName(self.panels_name).all([
        self._report.js.getElementsByName(self.tabs_name).all([
          self._report.js.data.all.element.css(css_not_clicked)]),
        tab.dom.css(dflt_css_clicked),
        self._report.js.data.all.element.hide(),
        tab_container.dom.toggleClass(css_cls_name, propagate=True) if css_cls_name is not None else "",
        div.dom.show()
      ])
    ])
    tab.inReport = False

    self.__panel_objs[name] = {"tab": tab_container, "content": div}
    if selected:
      # simulate a click on the tab
      pass
    return self

  def __str__(self):
    content = []
    for p in self.__panels:
      self.tabs_container += self.__panel_objs[p]["tab"]
      content.append(self.__panel_objs[p]["content"].html())
    return "<div %s>%s%s</div>%s" % (self.get_attrs(pyClassNames=self.style.get_classes()), self.tabs_container.html(), "".join(content), self.helper)


class IFrame(Html.Html):
  name, category, callFnc = 'IFrame', 'Container', 'iframe'

  def __init__(self, report, url, width, height, helper, profile):
    super(IFrame, self).__init__(report, url, css_attrs={"width": width, "height": height}, profile=profile)
    self.css({"overflow-x": 'hidden'})

  @property
  def _js__builder__(self):
    return 'htmlObj.src = data'

  def __str__(self):
    return "<iframe src='%s' %s frameborder='0' scrolling='no'></iframe>" % (self.val, self.get_attrs(pyClassNames=self.style.get_classes()))


class Dialog(Html.Html):
  name, category, callFnc = 'DialogMenu', 'Layouts', 'dialogs'
  __reqCss, __reqJs = ['jqueryui'], ['jqueryui']

  def __init__(self, report, recordSet, width, height, helper, profile):
    super(Dialog, self).__init__(report, recordSet,
                                 profile=profile)
    self.css({"border": '2px solid %s' % self._report.theme.greys[3], "display": "block", "position": "absolute",
              "background": self._report.theme.greys[0]})
    # self._report._props.setdefault('js', {}).setdefault("builders", []).append(self.dom.jquery_ui.draggable().toStr())

  @property
  def _js__builder__(self):
    return '''
      var div = jQuery("<div>")
      jQuery(htmlObj).append(div);
      div.dialog({modal: false, title: "rrrr"})'''

  # def jsAdd(self, title='data.event_val', isPyData=False):
  #   if isPyData:
  #     title = json.loads(title)
  #
  #   return '''
  #   var dialogWindow = $("<div>");
  #   var table = $("table");
  #   dialogWindow.append(table);
  #   dialogWindow.append('<input type="text">');
  #   var d = dialogWindow.dialog( { modal: false, title: %(title)s, show: 'puff', fluid: true,
  #       close: function () {$(this).remove()}, appendTo: "#%(htmlId)s", resizable: false,
  #       buttons: [{text: "Close", click: function() { $( this ).dialog("close")} } ]
  #   });
  #   d.parent().draggable({containment: '#%(htmlId)s'});
  #   event.preventDefault()''' % {'title': title, 'htmlId': self.htmlId}

  def __str__(self):
    return "<div %s></div>" % self.get_attrs(pyClassNames=self.defined)


class IconsMenu(Html.Html):
  name, category, callFnc = 'Icons Menu', 'Layouts', 'menu'
  __reqCss, __reqJs = ['font-awesome'], ['font-awesome']

  def __init__(self, icon_names, report, width, height, htmlCode, helper, profile):
    super(IconsMenu, self).__init__(report, None, width=width, css_attrs={"width": width, "height": height}, code=htmlCode,
                                    profile=profile)
    self._jsActions, self._definedActions = {}, []
    self._icons, self.icon = [], None
    self.css({"margin": "5px 0"})
    for i in icon_names:
      self.add_icon(i)

  def __getitem__(self, i):
    return self._icons[i]

  def add_icon(self, text, css=None, position="after"):
    """
    Add an icon to the HTML object

    Example
    checks.title.add_icon("fas fa-align-center")

    Documentation

    :param text: The icon reference from font awsome website
    :param css: Optional. A dictionary with the CSS style to be added to the component
    :param position:
    :return: The Html object
    """
    if text is not None:
      self._icons.append(self._report.ui.images.icon(text).css({"margin-right": '5px', 'cursor': "pointer"}))
      self.icon = self._icons[-1]
      if position == "before":
        self.prepend_child(self.icon)
      else:
        self.append_child(self.icon)
      #elf.icon.inReport = False
      if css is not None:
        self.icon.css(css)
    return self

  def add_select(self, action, data, width=150):
    options = ["<option>%s</option>" % d for d in data]
    self._jsActions[action] = '<select id="inputState" class="form-control" style="width:%spx;display:inline-block">%s</select>' % (width, "".join(options))
    self._definedActions.append(action)
    return self

  def __str__(self):
    htmlIcons = [htmlDef for action, htmlDef in self._jsActions.items()]
    return "<div %s>%s</div>" % (self.get_attrs(pyClassNames=self.style.get_classes()), "".join(htmlIcons))


class Form(Html.Html):
  name, category, callFnc = 'Generic Form', 'Forms', 'form'
  #_grpCls = CssGrpCls.CssGrpClassBase

  def __init__(self, report, htmlObjs, action, method, helper):
    super(Form, self).__init__(report, [])
    self.css({"padding": '5px'})
    if action is not None and method is not None:
      self.attr.update({"action": action, "method": method})
    self.add_helper(helper)
    self.submit = self._report.ui.button("Submit").set_attrs({"type": 'submit'})
    self.submit.inReport = False
    for i, htmlObj in enumerate(htmlObjs):
      self.__add__(htmlObj)

  def __add__(self, htmlObj):
    """ Add items to a container """
    htmlObj.inReport = False # Has to be defined here otherwise it is set too late
    self.val.append(htmlObj)
    return self

  def __str__(self):
    str_vals = "".join([i.html() for i in self.val]) if self.val is not None else ""
    return '<form %s>%s</form>%s' % (self.get_attrs(pyClassNames=self.defined), str_vals, self.helper)


class Modal(Html.Html):
  name, category, callFnc = 'Modal Popup',  'Container', 'modal'
  # _grpCls = CssGrpContainers.CssGrpClassModal

  def __init__(self, report, htmlObjs, submit, helper):
    super(Modal, self).__init__(report, [])
    self.add_helper(helper)
    self.doSubmit = submit
    if self.doSubmit:
      self.submit = report.ui.button("Submit").set_attrs({"type": 'submit'})
      self.submit.inReport = False
    self.col = report.ui.col([]).css({'border': '1px solid %s' % report.theme.greys[4],
                                      'width': 'auto', 'background-color': report.theme.greys[0]})
    self.closeBtn = report.ui.texts.span('&times', width='auto').css({'float': 'right', 'text-align': 'right',
                                                                      'margin-right': '10px', 'font-size': '24px',
                                                                      'z-index': 10})
    self.closeBtn.click(report.js.getElementById(self.htmlId).css({'display': "none"}))
    self.col += self.closeBtn
    self.col.inReport = False
    self.val.append(self.col)
    for htmlObj in htmlObjs:
      self.__add__(htmlObj)

  def __add__(self, htmlObj):
    """ Add items to a container """
    htmlObj.inReport = False # Has to be defined here otherwise it is set too late
    self.col += htmlObj
    return self

  def __str__(self):
    str_vals = "".join([i.html() for i in self.val]) if self.val is not None else ""
    if self.doSubmit:
      self.col += self.submit
    return '<div %s>%s</div>%s' % (self.get_attrs(pyClassNames=self.style.get_classes()), str_vals, self.helper)


class Indices(Html.Html):
  name, category, callFnc = 'Index', 'Panels', 'index'
  __reqCss, __reqJs = ['font-awesome'], ['font-awesome']

  def __init__(self, report, count, width, height, htmlCode, options, profile):
    super(Indices, self).__init__(report, count, css_attrs={"width": width, "height": height}, profile=profile)
    self.items = []
    self.__options = OptPanel.OptionsPanelPoints(report, options)
    for i in range(count):
      div = self._report.ui.div(i, width=(15, "px"))
      div.attr["name"] = self.htmlId
      div.attr["data-position"] = i + 1
      div.css({"display": 'inline-block', "padding": "2px", "text-align": "center"})
      div.css(self.options.div_css)
      div.style.add_classes.div.background_hover()
      div.inReport = False
      self.items.append(div)
    #
    self.first = self._report.ui.icon("fas fa-angle-double-left", width=(20, 'px')).css({"display": 'inline-block'})
    self.first.inReport = False
    self.prev = self._report.ui.icon("fas fa-chevron-left", width=(20, 'px')).css({"display": 'inline-block'})
    self.prev.inReport = False
    self.next = self._report.ui.icon("fas fa-chevron-right", width=(20, 'px')).css({"display": 'inline-block'})
    self.next.inReport = False
    self.last = self._report.ui.icon("fas fa-angle-double-right", width=(20, 'px')).css({"display": 'inline-block'})
    self.last.inReport = False

  @property
  def options(self):
    """

    :rtype: OptPanel.OptionsPanelPoints
    """
    return self.__options

  def __getitem__(self, i):
    return self.items[i]

  def click(self, i, jsFncs, profile=False):
    """

    :param i:
    :param jsFncs:
    :param profile:
    """
    if not isinstance(jsFncs, list):
      jsFncs = [jsFncs]
    return self[i].on("click", [
      self[i].dom.by_name.css({"border-bottom": "1px solid %s" % self._report.theme.colors[0]}).r,
      self[i].dom.css({"border-bottom": "1px solid %s" % self.options.background_color})] + jsFncs, profile)

  def __str__(self):
    str_vals = "".join([self.first.html(), self.prev.html()] + [i.html() for i in self.items] + [self.next.html(), self.last.html()])
    return '<div %s>%s</div>%s' % (self.get_attrs(pyClassNames=self.style.get_classes()), str_vals, self.helper)


class Points(Html.Html):
  name, category, callFnc = 'Index', 'Panels', 'index'

  def __init__(self, report, count, width, height, htmlCode, options, profile):
    super(Points, self).__init__(report, count, css_attrs={"width": width, "height": height}, profile=profile)
    self.items = []
    self.css({"text-align": "center"})
    self.__options = OptPanel.OptionsPanelPoints(report, options)
    for i in range(count):
      div = self._report.ui.div(self._report.entities.non_breaking_space)
      div.attr["name"] = self.htmlId
      div.attr["data-position"] = i + 1
      div.css({"border": "1px solid %s" % self._report.theme.greys[5], "border-radius": "10px", "width": "15px", "height": "15px"})
      div.css(self.options.div_css)
      div.style.add_classes.div.background_hover()
      div.inReport = False
      self.items.append(div)

  @property
  def options(self):
    """

    :rtype: OptPanel.OptionsPanelPoints
    """
    return self.__options

  def click(self, i, jsFncs, profile=False):
    """

    :param i:
    :param jsFncs:
    :param profile:
    """
    if not isinstance(jsFncs, list):
      jsFncs = [jsFncs]
    return self[i].on("click", [
      self[i].dom.by_name.css({"background-color": ""}).r,
      self[i].dom.css({"background-color": self.options.background_color})] + jsFncs, profile)

  def __getitem__(self, i):
    return self.items[i]

  def __str__(self):
    str_vals = "".join([i.html() for i in self.items])
    return '<div %s>%s</div>%s' % (self.get_attrs(pyClassNames=self.style.get_classes()), str_vals, self.helper)
